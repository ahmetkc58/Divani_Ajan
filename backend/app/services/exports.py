from pathlib import Path

import reportlab
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.schemas import DraftV1

WATERMARK = "SENTETİK TASLAK - RESMÎ BELGE DEĞİLDİR"


def export_docx(draft: DraftV1, output_path: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    marker = document.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = marker.add_run(WATERMARK)
    run.bold = True
    run.font.size = Pt(9)

    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(f"T.C.\n{_turkish_upper(draft.institution_name)}")
    header_run.bold = True
    header_run.font.size = Pt(12)

    metadata = document.add_paragraph()
    metadata.add_run(f"Sayı: {draft.number}").bold = True
    metadata.add_run(f"\tTarih: {draft.date}")

    subject = document.add_paragraph()
    subject.add_run(f"Konu: {draft.subject}").bold = True

    recipient = document.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recipient.add_run(_turkish_upper(draft.recipient_unit_name)).bold = True

    if draft.references:
        references = document.add_paragraph()
        references.add_run("İlgi: ").bold = True
        references.add_run("; ".join(draft.references))

    body = document.add_paragraph(draft.body)
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Cm(1.25)

    sign = document.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run(f"{draft.signatory}\nYetkili (Sentetik)").bold = True

    if draft.attachments:
        attachments = document.add_paragraph()
        attachments.add_run("Ekler:\n").bold = True
        attachments.add_run("\n".join(f"{index}. {item}" for index, item in enumerate(draft.attachments, 1)))
    if draft.distribution:
        distribution = document.add_paragraph()
        distribution.add_run("Dağıtım:\n").bold = True
        distribution.add_run("\n".join(draft.distribution))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(WATERMARK).italic = True
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def export_pdf(draft: DraftV1, output_path: Path) -> Path:
    font_dir = Path(reportlab.__file__).resolve().parent / "fonts"
    regular_font = font_dir / "Vera.ttf"
    bold_font = font_dir / "VeraBd.ttf"
    pdfmetrics.registerFont(TTFont("VeraTR", regular_font))
    pdfmetrics.registerFont(TTFont("VeraTR-Bold", bold_font))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.3 * cm,
        bottomMargin=2.3 * cm,
        title=draft.subject,
        author="Örnekşehir Belediyesi - Sentetik",
    )
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "BaseTR",
        parent=styles["Normal"],
        fontName="VeraTR",
        fontSize=10.5,
        leading=16,
        spaceAfter=10,
    )
    bold = ParagraphStyle("BoldTR", parent=base, fontName="VeraTR-Bold")
    centered = ParagraphStyle("CenteredTR", parent=bold, alignment=TA_CENTER)
    right = ParagraphStyle("RightTR", parent=bold, alignment=TA_RIGHT)
    body = ParagraphStyle("BodyTR", parent=base, alignment=TA_JUSTIFY, firstLineIndent=1.25 * cm)

    story = [
        Paragraph(WATERMARK, centered),
        Spacer(1, 0.2 * cm),
        Paragraph(f"T.C.<br/>{_escape(_turkish_upper(draft.institution_name))}", centered),
        Spacer(1, 0.4 * cm),
        Paragraph(f"<b>Sayı:</b> {_escape(draft.number)}&nbsp;&nbsp;&nbsp;&nbsp;<b>Tarih:</b> {_escape(draft.date)}", base),
        Paragraph(f"<b>Konu:</b> {_escape(draft.subject)}", base),
        Spacer(1, 0.4 * cm),
        Paragraph(_escape(_turkish_upper(draft.recipient_unit_name)), centered),
        Spacer(1, 0.5 * cm),
    ]
    if draft.references:
        story.append(Paragraph(f"<b>İlgi:</b> {_escape('; '.join(draft.references))}", base))
    story.extend(
        [
            Paragraph(_escape(draft.body).replace("\n", "<br/>"), body),
            Spacer(1, 0.5 * cm),
            Paragraph(f"{_escape(draft.signatory)}<br/>Yetkili (Sentetik)", right),
        ]
    )
    if draft.attachments:
        story.append(Paragraph("<b>Ekler:</b><br/>" + "<br/>".join(_escape(item) for item in draft.attachments), base))
    if draft.distribution:
        story.append(Paragraph("<b>Dağıtım:</b><br/>" + "<br/>".join(_escape(item) for item in draft.distribution), base))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(WATERMARK, centered))
    pdf.build(story)
    return output_path


def _escape(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _turkish_upper(value: str) -> str:
    return value.translate(str.maketrans({"i": "İ", "ı": "I"})).upper()
